"""
Document ingestion module for the RAG pipeline.
Supports multiple file formats and text chunking strategies.
"""
import os
import re
from pathlib import Path
from typing import List, Dict, Any, Optional, Union, BinaryIO
from dataclasses import dataclass
import json
import logging
from abc import ABC, abstractmethod

from pypdf import PdfReader
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

@dataclass
class DocumentChunk:
    """A chunk of text from a document with metadata."""
    text: str
    metadata: Dict[str, Any]
    chunk_id: Optional[str] = None

class DocumentLoader(ABC):
    """Base class for document loaders."""
    
    @abstractmethod
    def load(self, source: Union[str, Path, BinaryIO]) -> List[DocumentChunk]:
        """Load and chunk document from source."""
        pass

class TextLoader(DocumentLoader):
    """Load plain text files."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def load(self, source: Union[str, Path, BinaryIO]) -> List[DocumentChunk]:
        """Load and chunk text file."""
        if isinstance(source, (str, Path)):
            with open(source, 'r', encoding='utf-8') as f:
                text = f.read()
            file_metadata = {"source": str(source), "type": "text/plain"}
        else:
            text = source.read().decode('utf-8')
            file_metadata = {"source": "stream", "type": "text/plain"}
        
        return self._chunk_text(text, file_metadata)
    
    def _chunk_text(self, text: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Split text into overlapping chunks."""
        # Simple sentence-aware chunking
        sentences = re.split(r'(?<=[.!?])\s+', text)
        chunks = []
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            sentence_length = len(sentence.split())
            
            if current_chunk and current_length + sentence_length > self.chunk_size:
                # Add current chunk to chunks
                chunk_text = ' '.join(current_chunk)
                chunks.append(DocumentChunk(
                    text=chunk_text,
                    metadata=metadata.copy()
                ))
                
                # Start new chunk with overlap
                overlap_start = max(0, len(current_chunk) - self.chunk_overlap)
                current_chunk = current_chunk[overlap_start:]
                current_length = sum(len(s.split()) for s in current_chunk)
            
            current_chunk.append(sentence)
            current_length += sentence_length
        
        # Add the last chunk
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunks.append(DocumentChunk(
                text=chunk_text,
                metadata=metadata.copy()
           ))
        
        # Add chunk IDs
        for i, chunk in enumerate(chunks):
            chunk.chunk_id = f"chunk_{i:04d}"
        
        return chunks

class PDFLoader(DocumentLoader):
    """Load PDF files with pypdf."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_loader = TextLoader(chunk_size, chunk_overlap)
    
    def load(self, source: Union[str, Path, BinaryIO]) -> List[DocumentChunk]:
        """Load and chunk PDF file."""
        if isinstance(source, (str, Path)):
            with open(source, 'rb') as f:
                return self._load_pdf(f, str(source))
        return self._load_pdf(source, "stream")
    
    def _load_pdf(self, file_obj: BinaryIO, source: str) -> List[DocumentChunk]:
        """Extract text from PDF and chunk it."""
        try:
            pdf_reader = PdfReader(file_obj)
            text = ""
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += f"\n\n--- Page {page_num + 1} ---\n\n"
                text += page.extract_text() or ""
            
            metadata = {
                "source": source,
                "type": "application/pdf",
                "num_pages": len(pdf_reader.pages)
            }
            
            return self.text_loader._chunk_text(text, metadata)
            
        except Exception as e:
            logger.error(f"Error loading PDF {source}: {e}")
            return []

class DocxLoader(DocumentLoader):
    """Load DOCX files."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_loader = TextLoader(chunk_size, chunk_overlap)
    
    def load(self, source: Union[str, Path, BinaryIO]) -> List[DocumentChunk]:
        """Load and chunk DOCX file."""
        if isinstance(source, (str, Path)):
            doc = DocxDocument(source)
            source_str = str(source)
        else:
            doc = DocxDocument(source)
            source_str = "stream"
        
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        metadata = {
            "source": source_str,
            "type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        }
        
        return self.text_loader._chunk_text(text, metadata)

class JSONLoader(DocumentLoader):
    """Load JSON files."""
    
    def __init__(self, text_fields: List[str] = None):
        self.text_fields = text_fields or ["text", "content"]
    
    def load(self, source: Union[str, Path, BinaryIO]) -> List[DocumentChunk]:
        """Load JSON file and extract text from specified fields."""
        if isinstance(source, (str, Path)):
            with open(source, 'r', encoding='utf-8') as f:
                data = json.load(f)
            source_str = str(source)
        else:
            data = json.load(source)
            source_str = "stream"
        
        chunks = []
        
        # Handle list of objects
        if isinstance(data, list):
            for i, item in enumerate(data):
                text = self._extract_text(item)
                if text:
                    metadata = {
                        "source": f"{source_str}#item_{i}",
                        "type": "application/json"
                    }
                    chunks.append(DocumentChunk(
                        text=text,
                        metadata=metadata,
                        chunk_id=f"item_{i}"
                    ))
        # Handle single object
        elif isinstance(data, dict):
            text = self._extract_text(data)
            if text:
                metadata = {
                    "source": source_str,
                    "type": "application/json"
                }
                chunks.append(DocumentChunk(
                    text=text,
                    metadata=metadata,
                    chunk_id="root"
                ))
        
        return chunks
    
    def _extract_text(self, data: Any) -> str:
        """Extract text from JSON data."""
        if isinstance(data, str):
            return data
        elif isinstance(data, dict):
            # Try to find text fields
            for field in self.text_fields:
                if field in data and isinstance(data[field], str):
                    return data[field]
            # If no text field found, convert dict to string
            return json.dumps(data, ensure_ascii=False)
        elif isinstance(data, (list, tuple)):
            return "\n".join(str(item) for item in data)
        return str(data)

def get_loader(file_path: Union[str, Path]) -> Optional[DocumentLoader]:
    """Get appropriate loader for file type."""
    if isinstance(file_path, (str, Path)):
        ext = Path(file_path).suffix.lower()
    else:
        return None
    
    if ext == '.txt':
        return TextLoader()
    elif ext == '.pdf':
        return PDFLoader()
    elif ext == '.docx':
        return DocxLoader()
    elif ext == '.json':
        return JSONLoader()
    
    logger.warning(f"No loader found for file extension: {ext}")
    return None

def load_document(file_path: Union[str, Path]) -> List[DocumentChunk]:
    """Load document using appropriate loader."""
    loader = get_loader(file_path)
    if loader:
        return loader.load(file_path)
    return []

def load_directory(directory: Union[str, Path], 
                  extensions: List[str] = None) -> List[DocumentChunk]:
    """Load all documents from a directory with given extensions."""
    if extensions is None:
        extensions = ['.txt', '.pdf', '.docx', '.json']
    
    directory = Path(directory)
    if not directory.is_dir():
        raise ValueError(f"Not a directory: {directory}")
    
    all_chunks = []
    
    for ext in extensions:
        for file_path in directory.glob(f"**/*{ext}"):
            try:
                chunks = load_document(file_path)
                all_chunks.extend(chunks)
                logger.info(f"Loaded {len(chunks)} chunks from {file_path}")
            except Exception as e:
                logger.error(f"Error loading {file_path}: {e}")
    
    return all_chunks
